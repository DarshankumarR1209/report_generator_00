from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def calc_number_of_rows(number_of_images: int):
    num_of_cols_per_line = 2
    if number_of_images % num_of_cols_per_line == 0:
        return int(number_of_images / num_of_cols_per_line)
    else:
        return (number_of_images // num_of_cols_per_line) +1
    
def add_headers_footers(doc_path : str, doc_title : str, logo_path : str):
    report_doc = Document(doc_path)

    section = report_doc.sections[0]

    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(6))  # Adjust the width as needed

    # Add an image to the right side of the header
    image_path = logo_path  # Replace with the actual image path
    image_cell = header_table.cell(0, 1)
    image_paragraph = image_cell.paragraphs[0]
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Right-align the image
    image_run = image_paragraph.add_run()
    image_run.add_picture(image_path, width=Inches(0.5))  # Adjust the width as needed

    # Add the document title to the left side of the header
    title_cell = header_table.cell(0, 0)
    title_paragraph = title_cell.paragraphs[0]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Left-align the title
    title_run = title_paragraph.add_run(doc_title)
    title_run.font.size = Pt(12)  # Adjust font size as needed

    # Add a footer to the first page
    footer = section.footer

    # Add a paragraph to the footer for the page number
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center-align the page number

    # Create a page number field
    page_number_run = footer_paragraph.add_run()
    page_number_field = OxmlElement("w:fldSimple")
    page_number_field.set(qn("w:instr"), "PAGE")
    page_number_run._r.append(page_number_field)

    # Set the font size for the page number (optional)
    page_number_run.font.size = Pt(12)

    report_doc.save(doc_path)
