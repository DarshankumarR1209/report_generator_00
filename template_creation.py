from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import models

def ceate_reoprt(resized_img_dict : dict, no_of_imgs : int, no_of_rows : int ):

    # Create a new Document
    doc = Document()

    # Create page title
    pg_title = doc.add_heading("First Report", level=0)
    pg_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pg_title.bold= True

    # Add section title
    title = doc.add_heading("Images of damaged properties", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT  # left align the section title

    # Add content to the document
    paragraph = doc.add_paragraph("Below are the images of the damaged properties")

    # Add a table with necessary columns and rows
    no_of_cols = 0
    if no_of_imgs == 1:
        no_of_cols = 1
    elif no_of_imgs >1 :
        no_of_cols = 2
    # elif no_of_imgs > 2:
    #     no_of_cols = 3

    table = doc.add_table(rows=no_of_rows, cols=no_of_cols)
    table.autofit = True  # Disable autofit

    # Set column widths for the table
    # col_widths = (1, 1, 1)
    # for col_number, width in enumerate(col_widths):
    #     print(col_number,width)
    #     cell = table.cell(0, col_number)
    #     cell.width = Pt(width * 72)  # Convert inches to points (1 inch = 72 points)

    # # Add content to the table cells
    # table.cell(0, 0).text = "Column 1"
    # table.cell(0, 1).text = "Column 2"
    # table.cell(0, 2).text = "Column 3"

    # Add image to the cell
    col_number = 0
    for item_number, (file_name, file_path) in enumerate(resized_img_dict.items(), start=0):
        row_number = 0
        if item_number == 0 :
            row_number = 0
        elif item_number == 1:
            row_number = 0
        elif item_number > 1:
            row_number = item_number // no_of_cols
        # elif item_number > 2:
        #     row_number = item_number // no_of_cols
        print(f'Item Number is {item_number} and row number is {row_number} and column number is {col_number}')
        # for col_number in range(0,no_of_cols):
        
        cell = table.cell(row_number, col_number)
        paragraph = cell.add_paragraph()

        img_description = models.image2text(file_path)

        image_path = file_path
        run = paragraph.add_run()
        # run.add_picture(image_path, width = Pt(180), height = Pt(120))
        run.add_picture(image_path, height=Inches(2))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add file name
        text_para = cell.add_paragraph(img_description)
        text_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if col_number == no_of_cols - 1:
            col_number = 0
        else:
            col_number += 1

    # Save the DOCX file
    doc.save("report.docx")


def add_second_para(doc_path : str):
    report_doc = Document(doc_path)

    # Add section title
    title = report_doc.add_heading("Details of insurance and policy", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT  # left align the section title

    new_para = report_doc.add_paragraph('This is the following paragraph. It will have details of the client and policy details.')
    new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    report_doc.save(doc_path)